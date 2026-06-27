"""Tests for the DOCX generator module."""

from datetime import date
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

from app.document_generator import generate_docx


@pytest.fixture
def temp_dir(tmp_path):
    """Provide a temporary directory for test outputs."""
    return tmp_path


@pytest.fixture
def dummy_images(temp_dir):
    """Create a few small dummy images for testing."""
    images = []
    for i, color in enumerate([(255, 0, 0), (0, 255, 0), (0, 0, 255)]):
        img_path = temp_dir / f"dummy_{i}.png"
        img = Image.new("RGB", (100, 100), color)
        img.save(img_path)
        images.append(str(img_path))
    return images


@pytest.fixture
def project_details():
    """Sample project details as fetched from the project_details DB."""
    return {
        "region": "wilayah kerja UPT Gandul",
        "roadway": "jakarta - bandung",
        "tower_id": "Tower 123",
        "tower_type": "500kV",
    }


def _paragraph_texts(doc):
    return [p.text for p in doc.paragraphs]


def _all_texts(doc):
    """All visible text in the document: paragraphs + table cells."""
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return texts


def _specs_table(doc):
    """Return the project specs table (2-row, 4-col, first row has 'Penghantar')."""
    for table in doc.tables:
        if len(table.rows) == 2 and len(table.columns) == 4:
            if "Penghantar" in table.rows[0].cells[0].text:
                return table
    return None


def _footer_table(doc):
    """Return the footer table (1-row, 2-col, contains 'PT TESLA')."""
    for table in doc.tables:
        if len(table.rows) == 1 and len(table.columns) == 2:
            if "PT TESLA DAYA ELEKTRIKA" in table.rows[0].cells[0].text:
                return table
    return None


class TestGenerateDocx:
    """Tests for the generate_docx function."""

    def test_docx_created(self, temp_dir, project_details, dummy_images):
        """A .docx file is created at the specified output path."""
        output_path = temp_dir / "report.docx"
        generate_docx(
            project_details=project_details,
            report_date=date(2026, 6, 10),
            grouped_data={"Section atas": [dummy_images[0]]},
            measurement_images=[],
            output_path=str(output_path),
        )
        assert output_path.exists()
        assert output_path.stat().st_size > 0

    def test_title_block_present(self, temp_dir, project_details, dummy_images):
        """The centered bold title and subtitle are present."""
        output_path = temp_dir / "report.docx"
        generate_docx(
            project_details=project_details,
            report_date=date(2026, 6, 10),
            grouped_data={"Section atas": [dummy_images[0]]},
            measurement_images=[],
            output_path=str(output_path),
        )
        doc = Document(str(output_path))
        texts = _paragraph_texts(doc)
        assert "LAPORAN DOKUMENTASI PEMASANGAN" in texts
        assert "PEKERJAAN PENGADAAN DAN PEMASANGAN PROTEKSI PETIR DI wilayah kerja UPT Gandul" in texts
        title_para = doc.paragraphs[0]
        assert title_para.alignment is not None  # centered
        assert title_para.runs[0].bold is True

    def test_project_specs_present(self, temp_dir, project_details, dummy_images):
        """Project specs (roadway, report_date, tower_id, tower_type) are embedded."""
        output_path = temp_dir / "report.docx"
        generate_docx(
            project_details=project_details,
            report_date=date(2026, 4, 16),
            grouped_data={"Section atas": [dummy_images[0]]},
            measurement_images=[],
            output_path=str(output_path),
        )
        doc = Document(str(output_path))
        body = "\n".join(_all_texts(doc))
        assert "jakarta - bandung" in body
        assert "16 April 2026" in body
        assert "Tower 123" in body
        assert "500kV" in body

    def test_project_specs_table_structure(self, temp_dir, project_details, dummy_images):
        """Specs are rendered as a 2-row, 4-col borderless table with aligned labels."""
        output_path = temp_dir / "report.docx"
        generate_docx(
            project_details=project_details,
            report_date=date(2026, 4, 16),
            grouped_data={"Section atas": [dummy_images[0]]},
            measurement_images=[],
            output_path=str(output_path),
        )
        doc = Document(str(output_path))
        specs = _specs_table(doc)
        assert specs is not None, "Specs table not found"
        row0 = [c.text for c in specs.rows[0].cells]
        row1 = [c.text for c in specs.rows[1].cells]
        assert row0 == ["Penghantar", ": jakarta - bandung", "Tanggal Pemasangan", ": 16 April 2026"]
        assert row1 == ["No Tower", ": Tower 123", "Tipe tower", ": 500kV"]

    def test_progress_pekerjaan_heading(self, temp_dir, project_details, dummy_images):
        """Progress Pekerjaan heading is present."""
        output_path = temp_dir / "report.docx"
        generate_docx(
            project_details=project_details,
            report_date=date(2026, 6, 10),
            grouped_data={"Section atas": [dummy_images[0]]},
            measurement_images=[],
            output_path=str(output_path),
        )
        doc = Document(str(output_path))
        assert "Progress Pekerjaan" in _paragraph_texts(doc)

    def test_section_headings_only_when_images_exist(self, temp_dir, project_details, dummy_images):
        """Section sub-headings appear only for sections that have images."""
        output_path = temp_dir / "report.docx"
        generate_docx(
            project_details=project_details,
            report_date=date(2026, 6, 10),
            grouped_data={
                "Section atas": [dummy_images[0]],
                "Section bawah": [dummy_images[1]],
            },
            measurement_images=[],
            output_path=str(output_path),
        )
        doc = Document(str(output_path))
        texts = _paragraph_texts(doc)
        assert "Section Atas" in texts
        assert "Section Bawah" in texts
        assert "Section Tengah" not in texts

    def test_section_ordering(self, temp_dir, project_details, dummy_images):
        """Sections are ordered Atas -> Tengah -> Bawah regardless of input order."""
        output_path = temp_dir / "report.docx"
        generate_docx(
            project_details=project_details,
            report_date=date(2026, 6, 10),
            grouped_data={
                "Section bawah": [dummy_images[2]],
                "Section atas": [dummy_images[0]],
                "Section tengah": [dummy_images[1]],
            },
            measurement_images=[],
            output_path=str(output_path),
        )
        doc = Document(str(output_path))
        texts = _paragraph_texts(doc)
        assert texts.index("Section Atas") < texts.index("Section Tengah")
        assert texts.index("Section Tengah") < texts.index("Section Bawah")

    def test_english_section_aliases(self, temp_dir, project_details, dummy_images):
        """English section aliases (top/mid/bottom) map to Atas/Tengah/Bawah."""
        output_path = temp_dir / "report.docx"
        generate_docx(
            project_details=project_details,
            report_date=date(2026, 6, 10),
            grouped_data={
                "Section top": [dummy_images[0]],
                "Section bottom": [dummy_images[1]],
            },
            measurement_images=[],
            output_path=str(output_path),
        )
        doc = Document(str(output_path))
        texts = _paragraph_texts(doc)
        assert "Section Atas" in texts
        assert "Section Bawah" in texts

    def test_measurement_section_absent_when_no_alat_ukur(self, temp_dir, project_details, dummy_images):
        """Progress pengukuran heading is absent when no measurement images."""
        output_path = temp_dir / "report.docx"
        generate_docx(
            project_details=project_details,
            report_date=date(2026, 6, 10),
            grouped_data={"Section atas": [dummy_images[0]]},
            measurement_images=[],
            output_path=str(output_path),
        )
        doc = Document(str(output_path))
        assert "Progress pengukuran" not in _paragraph_texts(doc)

    def test_measurement_section_present_when_alat_ukur(self, temp_dir, project_details, dummy_images):
        """Progress pengukuran heading is present when measurement images exist."""
        output_path = temp_dir / "report.docx"
        generate_docx(
            project_details=project_details,
            report_date=date(2026, 6, 10),
            grouped_data={"Section atas": [dummy_images[0]]},
            measurement_images=[dummy_images[1], dummy_images[2]],
            output_path=str(output_path),
        )
        doc = Document(str(output_path))
        assert "Progress pengukuran" in _paragraph_texts(doc)

    def test_images_embedded(self, temp_dir, project_details, dummy_images):
        """Images are embedded as inline shapes in the document."""
        output_path = temp_dir / "report.docx"
        generate_docx(
            project_details=project_details,
            report_date=date(2026, 6, 10),
            grouped_data={"Section atas": dummy_images},
            measurement_images=[],
            output_path=str(output_path),
        )
        doc = Document(str(output_path))
        assert len(doc.inline_shapes) == 3

    def test_footer_present(self, temp_dir, project_details, dummy_images):
        """Footer with company names is present in a single-row borderless table."""
        output_path = temp_dir / "report.docx"
        generate_docx(
            project_details=project_details,
            report_date=date(2026, 6, 10),
            grouped_data={"Section atas": [dummy_images[0]]},
            measurement_images=[],
            output_path=str(output_path),
        )
        doc = Document(str(output_path))
        footer = _footer_table(doc)
        assert footer is not None, "Footer table not found"
        assert footer.rows[0].cells[0].text == "PT TESLA DAYA ELEKTRIKA"
        assert footer.rows[0].cells[1].text == "PT PLN (PERSERO)"
        # Verify bold
        assert footer.rows[0].cells[0].paragraphs[0].runs[0].bold is True
        assert footer.rows[0].cells[1].paragraphs[0].runs[0].bold is True

    def test_empty_sections(self, temp_dir, project_details):
        """Even with no images, a .docx with title and footer is generated."""
        output_path = temp_dir / "empty.docx"
        generate_docx(
            project_details=project_details,
            report_date=date(2026, 6, 10),
            grouped_data={},
            measurement_images=[],
            output_path=str(output_path),
        )
        assert output_path.exists()
        doc = Document(str(output_path))
        texts = _paragraph_texts(doc)
        assert "LAPORAN DOKUMENTASI PEMASANGAN" in texts
        assert "Progress Pekerjaan" in texts
        assert "Progress pengukuran" not in texts