"""Document generator module for creating report .docx files with image grids."""

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches

from PIL import Image


# Image grid layout constants
GRID_COLS = 3
# Maximum height (in inches) for a single image cell to prevent overly tall images.
MAX_CELL_HEIGHT_IN = 2.2
# Width (in inches) of each image cell in the 3-column grid.
CELL_WIDTH_IN = 2.0

# Project specs table column widths (in inches).
_SPECS_COL_WIDTHS = [1.0, 2.2, 1.6, 1.4]

# Fixed display order of sections in the Progress Pekerjaan block.
# Maps the position keyword (lowercase, normalized from sub_id) to the display name.
_SECTION_ORDER = ["atas", "tengah", "bawah"]
_SECTION_ALIASES = {
    "top": "atas",
    "mid": "tengah",
    "bottom": "bawah",
    "atas": "atas",
    "tengah": "tengah",
    "bawah": "bawah",
}


def _section_position(sub_id: str) -> str | None:
    """Extract the normalized section position keyword from a sub_id string."""
    match = re.search(r"section\s*[:\-]?\s*([a-zA-Z]+)", sub_id or "", re.IGNORECASE)
    if not match:
        return None
    word = match.group(1).lower()
    return _SECTION_ALIASES.get(word)


def _format_report_date(report_date: date) -> str:
    """Format a report date as 'DD Month YYYY' (Indonesian month names)."""
    months_id = [
        "Januari", "Februari", "Maret", "April", "Mei", "Juni",
        "Juli", "Agustus", "September", "Oktober", "November", "Desember",
    ]
    return f"{report_date.day} {months_id[report_date.month - 1]} {report_date.year}"


def _capitalize_value(value: str) -> str:
    """
    Capitalize the first letter of each whitespace-separated token,
    preserving the remaining characters (e.g. '500kV' stays '500kV').

    Tokens that start with a non-letter (digits, symbols) are left unchanged.
    """
    tokens = []
    for token in value.split():
        if token and token[0].isalpha():
            tokens.append(token[0].upper() + token[1:])
        else:
            tokens.append(token)
    return " ".join(tokens)


def _add_centered_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    """Add a centered paragraph with optional bold text."""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.bold = bold
    return para


def _add_image_grid(doc: Document, image_paths: list[str]) -> None:
    """Render images in a 3-column borderless table, scaling to fit cell width."""
    if not image_paths:
        return

    # Build rows of GRID_COLS images each.
    rows = []
    current = []
    for img_path in image_paths:
        if not Path(img_path).exists():
            continue
        current.append(img_path)
        if len(current) >= GRID_COLS:
            rows.append(current)
            current = []
    if current:
        rows.append(current)

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=GRID_COLS)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _remove_table_borders(table)

    for r, row_images in enumerate(rows):
        for c in range(GRID_COLS):
            cell = table.rows[r].cells[c]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if c < len(row_images):
                _add_scaled_image(cell, row_images[c])


def _remove_table_borders(table) -> None:
    """Remove all borders from a table by setting its borders to nil."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    # Remove existing tblBorders if present
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "none",
            qn("w:sz"): "0",
            qn("w:space"): "0",
            qn("w:color"): "auto",
        })
        borders.append(el)
    tblPr.append(borders)


def _add_scaled_image(cell, img_path: str) -> None:
    """Add an image to a table cell, scaled to CELL_WIDTH_IN with capped height."""
    with Image.open(img_path) as img:
        orig_w, orig_h = img.size

    width = CELL_WIDTH_IN
    scale_h = width * orig_h / orig_w
    height = min(scale_h, MAX_CELL_HEIGHT_IN)

    para = cell.paragraphs[0]
    run = para.add_run()
    run.add_picture(img_path, width=Inches(width), height=Inches(height))


def _set_cell_text(cell, text: str, bold: bool = False,
                   align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Replace a cell's paragraph text with a single bold run."""
    para = cell.paragraphs[0]
    # Clear any existing runs
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold


def _set_col_widths(table, widths: list[float]) -> None:
    """Set fixed column widths (in inches) on every cell for consistent layout."""
    from docx.shared import Inches as _Inches
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = _Inches(w)


def _add_specs_table(
    doc: Document,
    roadway: str,
    date_text: str,
    tower_id: str,
    tower_type: str,
) -> None:
    """Add the project specs as a 2-row, 4-column borderless table."""
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _remove_table_borders(table)
    _set_col_widths(table, _SPECS_COL_WIDTHS)

    rows = [
        ("Penghantar", f": {roadway}", "Tanggal Pemasangan", f": {date_text}"),
        ("No Tower", f": {tower_id}", "Tipe tower", f": {tower_type}"),
    ]
    for r, (label1, value1, label2, value2) in enumerate(rows):
        _set_cell_text(table.rows[r].cells[0], label1)
        _set_cell_text(table.rows[r].cells[1], value1)
        _set_cell_text(table.rows[r].cells[2], label2)
        _set_cell_text(table.rows[r].cells[3], value2)


def _add_footer_table(doc: Document) -> None:
    """Add the two-company footer as a 1-row, 2-column borderless table (one line)."""
    table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(table)
    _set_col_widths(table, [3.1, 3.1])

    _set_cell_text(
        table.rows[0].cells[0], "PT TESLA DAYA ELEKTRIKA",
        bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    _set_cell_text(
        table.rows[0].cells[1], "PT PLN (PERSERO)",
        bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
    )


def generate_docx(
    project_details: dict | None,
    report_date: date,
    grouped_data: dict,
    measurement_images: list[str],
    output_path: str,
):
    """
    Generate a .docx report matching the contoh-progress-report layout.

    Parameters
    ----------
    project_details : dict | None
        Project details dict with keys: region, roadway, tower_id, tower_type,
        project_name. If None, the report is still generated with blank placeholders.
    report_date : date
        The report date (from photo metadata). Displayed as 'Tanggal Pemasangan'.
    grouped_data : dict
        Mapping of sub_id (e.g. 'Section atas') -> list of local image file paths
        for the Progress Pekerjaan section.
    measurement_images : list[str]
        List of local image file paths tagged as 'Alat Ukur'. Rendered only in
        the Progress pengukuran section. If empty, that section is omitted.
    output_path : str
        Where to write the resulting .docx file.
    """
    details = project_details or {}
    region = details.get("region", "") or ""
    roadway = _capitalize_value(details.get("roadway", "") or "")
    tower_id = _capitalize_value(details.get("tower_id", "") or "")
    tower_type = _capitalize_value(details.get("tower_type", "") or "")
    project_name = details.get("project_name", "") or ""
    date_text = _format_report_date(report_date)

    doc = Document()

    # ── Title block ──
    _add_centered_paragraph(doc, "LAPORAN DOKUMENTASI PEMASANGAN PEKERJAAN", bold=True)
    _add_centered_paragraph(
        doc,
        f"{project_name} DI WILAYAH KERJA {region}",
        bold=True,
    )
    doc.add_paragraph()

    # ── Project specs (borderless table for vertical alignment) ──
    _add_specs_table(doc, roadway, date_text, tower_id, tower_type)
    doc.add_paragraph()

    # ── Progress Pekerjaan ──
    _add_centered_paragraph(doc, "Progress Pekerjaan")

    # Order sections by fixed order, only rendering those with images.
    buckets: dict[str, list[str]] = {}
    for sub_id, images in grouped_data.items():
        pos = _section_position(sub_id) or sub_id
        buckets.setdefault(pos, [])
        buckets[pos].extend(images)

    for pos in _SECTION_ORDER:
        images = buckets.get(pos)
        if not images:
            continue
        display_name = f"Section {pos.capitalize()}"
        _add_centered_paragraph(doc, display_name)
        _add_image_grid(doc, images)

    # ── Progress pengukuran (optional) ──
    if measurement_images:
        _add_centered_paragraph(doc, "Progress pengukuran")
        _add_image_grid(doc, measurement_images)

    # Spacer before footer
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # ── Footer (borderless table, guaranteed single line) ──
    _add_footer_table(doc)

    doc.save(output_path)
    return output_path